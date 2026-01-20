"""
TemplateAnalysisToolkit - Extracts document structure from sample files.

This toolkit allows the Document Agent to analyze sample documents
and extract their structural outline (headings, sections, tables)
to replicate the format in generated documents.
"""

import os
import re
from typing import List, Dict, Any
from camel.toolkits.function_tool import FunctionTool
from app.utils.toolkit.abstract_toolkit import AbstractToolkit
from app.utils.listen.toolkit_listen import listen_toolkit
from app.service.task import Agents
from utils import traceroot_wrapper as traceroot

logger = traceroot.get_logger("template_analysis_toolkit")


class TemplateAnalysisToolkit(AbstractToolkit):
    """Toolkit for analyzing document structure from sample files."""
    
    agent_name: str = Agents.document_agent

    def __init__(self, api_task_id: str, working_directory: str | None = None):
        self.api_task_id = api_task_id
        if working_directory is None:
            from app.component.environment import env
            working_directory = env("file_save_path", os.path.expanduser("~/Downloads"))
        self.working_directory = working_directory

    @listen_toolkit(
        None,
        lambda self, file_path: f"Analyzing document structure: {os.path.basename(file_path)}"
    )
    def analyze_document_structure(self, file_path: str) -> str:
        """
        Analyzes a sample document and extracts its structural outline.
        
        This tool reads a document file and extracts:
        - Headings and their hierarchy (H1, H2, H3, etc.)
        - Section names and their nesting level
        - Table structures (column headers, approximate row counts)
        - List structures
        
        Use this tool BEFORE writing any document when a sample/template 
        file is provided. The extracted structure should be replicated 
        in the generated document.
        
        Args:
            file_path (str): Absolute path to the sample document file.
                            Supports: .docx, .md, .txt, .html
        
        Returns:
            str: A structured outline of the document in Markdown format,
                 showing headings, sections, and table structures.
        """
        if not os.path.exists(file_path):
            return f"Error: File not found at {file_path}"
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_ext == '.docx':
                return self._analyze_docx(file_path)
            elif file_ext == '.md':
                return self._analyze_markdown(file_path)
            elif file_ext == '.txt':
                return self._analyze_text(file_path)
            elif file_ext == '.html':
                return self._analyze_html(file_path)
            elif file_ext == '.pdf':
                return self._analyze_pdf(file_path)
            else:
                return f"Unsupported file format: {file_ext}. Supported: .docx, .md, .txt, .html, .pdf"
        except Exception as e:
            logger.error(f"Error analyzing document: {e}", exc_info=True)
            return f"Error analyzing document: {str(e)}"
    
    def _analyze_pdf(self, file_path: str) -> str:
        """Extract structure from PDF files."""
        try:
            from pypdf import PdfReader
        except ImportError:
            return "Error: pypdf library not installed. Install with: pip install pypdf"
        
        try:
            reader = PdfReader(file_path)
            full_text = ""
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    full_text += text + "\n"
            
            # Use text analysis on extracted content
            # Add specific PDF metadata if available
            outline = ["# Document Structure Analysis (PDF)", ""]
            outline.append(f"**Source File:** {os.path.basename(file_path)}")
            outline.append(f"**Pages:** {len(reader.pages)}")
            
            # Reuse logic from _analyze_text but tailored
            # We call a helper that returns list of lines, then join
            text_analysis = self._analyze_text_content(full_text)
            outline.append(text_analysis)
            
            return "\n".join(outline)
        except Exception as e:
            return f"Error reading PDF: {str(e)}"

    def _analyze_text_content(self, content: str) -> str:
        """Helper to analyze structure from raw text string."""
        outline = []
        outline.append("")
        outline.append("## Detected Sections (Heuristic)")
        outline.append("")
        
        lines = content.split('\n')
        section_count = 0
        
        for line in lines:
            text = line.strip()
            if not text:
                continue
            
            # Heuristics for section headers
            is_header = False
            level = 2
            
            if text.isupper() and len(text) < 80 and len(text) > 3:
                is_header = True
                level = 1
            elif re.match(r'^[IVX]+\.\s+', text):
                is_header = True
                level = 1
            elif re.match(r'^[A-Z]\.\s+', text):
                is_header = True
                level = 2
            elif re.match(r'^\d+\.\s+[A-Z]', text) and len(text) < 100:
                is_header = True
                level = 2
            elif re.match(r'^\d+\.\d+\.\s+', text):
                is_header = True
                level = 3
            elif text.endswith(':') and len(text) < 60:
                is_header = True
                level = 2
            
            if is_header:
                section_count += 1
                indent = "  " * (level - 1)
                display_text = text[:60] + "..." if len(text) > 60 else text
                outline.append(f"{indent}- **[Level {level}]** {display_text}")
        
        outline.append("")
        outline.append("## Summary")
        outline.append(f"- Detected sections: {section_count}")
        outline.append("")
        outline.append("> **INSTRUCTION:** Replicate this structure in your generated document.")
        
        return "\n".join(outline)
    
    def _analyze_docx(self, file_path: str) -> str:
        """Extract structure from Word documents."""

        try:
            from docx import Document
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            return "Error: python-docx library not installed. Install with: pip install python-docx"
        
        doc = Document(file_path)
        outline = ["# Document Structure Analysis", ""]
        outline.append(f"**Source File:** {os.path.basename(file_path)}")
        outline.append("")
        outline.append("## Sections and Headings")
        outline.append("")
        
        current_level = 0
        section_count = 0
        table_count = 0
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            style_name = para.style.name.lower() if para.style else ""
            
            # Detect headings by style or formatting
            level = 0
            if 'heading 1' in style_name or 'title' in style_name:
                level = 1
            elif 'heading 2' in style_name:
                level = 2
            elif 'heading 3' in style_name:
                level = 3
            elif 'heading 4' in style_name:
                level = 4
            elif para.runs and para.runs[0].bold and len(text) < 100:
                # Bold short text is likely a heading
                level = 2 if current_level < 2 else current_level + 1
            
            if level > 0:
                section_count += 1
                indent = "  " * (level - 1)
                outline.append(f"{indent}- **[Level {level}]** {text}")
                current_level = level
        
        # Analyze tables
        if doc.tables:
            outline.append("")
            outline.append("## Table Structures")
            outline.append("")
            for i, table in enumerate(doc.tables):
                table_count += 1
                rows = len(table.rows)
                cols = len(table.columns) if table.rows else 0
                
                # Get header row if exists
                headers = []
                if table.rows:
                    for cell in table.rows[0].cells:
                        cell_text = cell.text.strip()[:30]
                        if cell_text:
                            headers.append(cell_text)
                
                outline.append(f"### Table {i+1}")
                outline.append(f"- **Rows:** {rows}")
                outline.append(f"- **Columns:** {cols}")
                if headers:
                    outline.append(f"- **Headers:** {', '.join(headers)}")
                outline.append("")
        
        # Summary
        outline.append("## Summary")
        outline.append(f"- Total sections/headings: {section_count}")
        outline.append(f"- Total tables: {table_count}")
        outline.append("")
        outline.append("> **INSTRUCTION:** Your generated document MUST replicate this structure.")
        outline.append("> Each section listed above should appear in your output with equivalent content depth.")
        
        return "\n".join(outline)

    def _analyze_markdown(self, file_path: str) -> str:
        """Extract structure from Markdown files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        outline = ["# Document Structure Analysis", ""]
        outline.append(f"**Source File:** {os.path.basename(file_path)}")
        outline.append("")
        outline.append("## Sections and Headings")
        outline.append("")
        
        lines = content.split('\n')
        section_count = 0
        table_count = 0
        in_table = False
        current_table_headers = []
        current_table_rows = 0
        
        for line in lines:
            # Detect headings
            heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
            if heading_match:
                section_count += 1
                level = len(heading_match.group(1))
                text = heading_match.group(2)
                indent = "  " * (level - 1)
                outline.append(f"{indent}- **[Level {level}]** {text}")
            
            # Detect tables
            if '|' in line and line.strip().startswith('|'):
                if not in_table:
                    in_table = True
                    table_count += 1
                    current_table_rows = 0
                    # Parse headers
                    cells = [c.strip() for c in line.split('|') if c.strip()]
                    current_table_headers = cells
                else:
                    if not line.replace('|', '').replace('-', '').replace(':', '').strip():
                        continue  # Skip separator line
                    current_table_rows += 1
            else:
                if in_table:
                    # End of table, record it
                    outline.append("")
                    outline.append(f"### Table {table_count}")
                    outline.append(f"- **Rows:** ~{current_table_rows}")
                    if current_table_headers:
                        outline.append(f"- **Headers:** {', '.join(current_table_headers[:5])}")
                    in_table = False
        
        # Close any remaining table
        if in_table:
            outline.append("")
            outline.append(f"### Table {table_count}")
            outline.append(f"- **Rows:** ~{current_table_rows}")
            if current_table_headers:
                outline.append(f"- **Headers:** {', '.join(current_table_headers[:5])}")
        
        outline.append("")
        outline.append("## Summary")
        outline.append(f"- Total sections/headings: {section_count}")
        outline.append(f"- Total tables: {table_count}")
        outline.append("")
        outline.append("> **INSTRUCTION:** Your generated document MUST replicate this structure.")
        
        return "\n".join(outline)

    def _analyze_text(self, file_path: str) -> str:
        """Extract structure from plain text files using heuristics."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        outline = ["# Document Structure Analysis", ""]
        outline.append(f"**Source File:** {os.path.basename(file_path)}")
        outline.append("")
        outline.append("## Detected Sections (Heuristic)")
        outline.append("")
        
        lines = content.split('\n')
        section_count = 0
        
        for line in lines:
            text = line.strip()
            if not text:
                continue
            
            # Heuristics for section headers:
            # - All caps
            # - Numbered (1., I., A.)
            # - Short lines with colon
            is_header = False
            level = 2
            
            if text.isupper() and len(text) < 80 and len(text) > 3:
                is_header = True
                level = 1
            elif re.match(r'^[IVX]+\.\s+', text):
                is_header = True
                level = 1
            elif re.match(r'^[A-Z]\.\s+', text):
                is_header = True
                level = 2
            elif re.match(r'^\d+\.\s+[A-Z]', text) and len(text) < 100:
                is_header = True
                level = 2
            elif re.match(r'^\d+\.\d+\.\s+', text):
                is_header = True
                level = 3
            elif text.endswith(':') and len(text) < 60:
                is_header = True
                level = 2
            
            if is_header:
                section_count += 1
                indent = "  " * (level - 1)
                display_text = text[:60] + "..." if len(text) > 60 else text
                outline.append(f"{indent}- **[Level {level}]** {display_text}")
        
        outline.append("")
        outline.append("## Summary")
        outline.append(f"- Detected sections: {section_count}")
        outline.append("")
        outline.append("> **INSTRUCTION:** Replicate this structure in your generated document.")
        
        return "\n".join(outline)

    def _analyze_html(self, file_path: str) -> str:
        """Extract structure from HTML files."""
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        outline = ["# Document Structure Analysis", ""]
        outline.append(f"**Source File:** {os.path.basename(file_path)}")
        outline.append("")
        outline.append("## Sections and Headings")
        outline.append("")
        
        section_count = 0
        
        # Simple regex-based extraction for h1-h6 tags
        for level in range(1, 7):
            pattern = rf'<h{level}[^>]*>(.*?)</h{level}>'
            matches = re.findall(pattern, content, re.IGNORECASE | re.DOTALL)
            for match in matches:
                section_count += 1
                text = re.sub(r'<[^>]+>', '', match).strip()[:60]
                indent = "  " * (level - 1)
                outline.append(f"{indent}- **[Level {level}]** {text}")
        
        # Count tables
        table_count = len(re.findall(r'<table', content, re.IGNORECASE))
        
        outline.append("")
        outline.append("## Summary")
        outline.append(f"- Total sections: {section_count}")
        outline.append(f"- Total tables: {table_count}")
        outline.append("")
        outline.append("> **INSTRUCTION:** Replicate this structure in your generated document.")
        
        return "\n".join(outline)

    @listen_toolkit(
        None,
        lambda self, file_path: f"Analyzing writing style: {os.path.basename(file_path)}"
    )
    def analyze_writing_style(self, file_path: str) -> str:
        """
        Analyzes a sample document and extracts its writing style characteristics.
        
        This tool examines a template document to understand:
        - Formality level (formal, semi-formal, casual)
        - Sentence structure patterns (complex, simple, mixed)
        - Vocabulary type (technical, academic, conversational)
        - Common phrases and sentence starters used
        - Tone (professional, persuasive, informative, etc.)
        
        Use this tool on the TEMPLATE file (File 1) to understand the writing 
        style that should be replicated when writing new content based on 
        reference materials (File 2).
        
        Args:
            file_path (str): Absolute path to the template document file.
                            Supports: .docx, .md, .txt
        
        Returns:
            str: A style guide describing the writing characteristics to replicate.
        """
        if not os.path.exists(file_path):
            return f"Error: File not found at {file_path}"
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            # Read content based on file type
            if file_ext == '.docx':
                try:
                    from docx import Document
                    doc = Document(file_path)
                    content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                except ImportError:
                    return "Error: python-docx library not installed"
            elif file_ext in ['.md', '.txt']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
            else:
                return f"Unsupported file format for style analysis: {file_ext}"
            
            return self._analyze_style(content, os.path.basename(file_path))
        except Exception as e:
            logger.error(f"Error analyzing writing style: {e}", exc_info=True)
            return f"Error analyzing writing style: {str(e)}"

    def _analyze_style(self, content: str, filename: str) -> str:
        """Analyze writing style from text content."""
        # Clean content
        lines = [l.strip() for l in content.split('\n') if l.strip()]
        sentences = []
        for line in lines:
            # Skip headers (lines that are all caps or start with #)
            if line.isupper() or line.startswith('#'):
                continue
            # Split into sentences
            for sent in re.split(r'[.!?]+', line):
                sent = sent.strip()
                if len(sent) > 20:  # Only consider meaningful sentences
                    sentences.append(sent)
        
        if not sentences:
            return "Error: Not enough text content to analyze style."
        
        # Extract sample sentences (first 5 and last 5)
        sample_sentences = sentences[:5] + sentences[-5:] if len(sentences) > 10 else sentences
        
        # Analyze formality indicators
        formal_indicators = ['theo đó', 'nhằm mục đích', 'được thực hiện', 'căn cứ', 
                           'đề xuất', 'kiến nghị', 'trình bày', 'nghiên cứu',
                           'therefore', 'consequently', 'furthermore', 'moreover',
                           'in accordance with', 'pursuant to', 'hereby']
        
        casual_indicators = ['như này', 'thế này', 'vậy đó', 'được rồi',
                           'gonna', 'wanna', 'kinda', 'stuff', 'things',
                           'ok', 'bạn', 'mình']
        
        formal_count = sum(1 for sent in sentences for ind in formal_indicators if ind.lower() in sent.lower())
        casual_count = sum(1 for sent in sentences for ind in casual_indicators if ind.lower() in sent.lower())
        
        # Determine formality
        if formal_count > casual_count * 2:
            formality = "Rất trang trọng (Very Formal) - Văn phong hành chính, học thuật"
        elif formal_count > casual_count:
            formality = "Trang trọng (Formal) - Văn phong chuyên nghiệp, lịch sự"
        elif casual_count > formal_count:
            formality = "Thân mật (Casual) - Văn phong gần gũi, dễ tiếp cận"
        else:
            formality = "Trung tính (Neutral) - Văn phong cân bằng"
        
        # Calculate average sentence length
        avg_sentence_length = sum(len(s.split()) for s in sentences) / len(sentences) if sentences else 0
        
        if avg_sentence_length > 25:
            sentence_style = "Câu dài, phức tạp - Nhiều mệnh đề, giải thích chi tiết"
        elif avg_sentence_length > 15:
            sentence_style = "Câu trung bình - Cân bằng giữa chi tiết và dễ đọc"
        else:
            sentence_style = "Câu ngắn gọn - Súc tích, dễ hiểu"
        
        # Extract common sentence starters
        starters = {}
        for sent in sentences:
            words = sent.split()
            if len(words) >= 2:
                starter = ' '.join(words[:2])
                starters[starter] = starters.get(starter, 0) + 1
        
        common_starters = sorted(starters.items(), key=lambda x: -x[1])[:5]
        
        # Build style guide
        style_guide = ["# 📝 Writing Style Analysis / Phân Tích Giọng Văn", ""]
        style_guide.append(f"**Source File:** {filename}")
        style_guide.append("")
        
        style_guide.append("## Formality Level / Độ Trang Trọng")
        style_guide.append(f"**{formality}**")
        style_guide.append("")
        
        style_guide.append("## Sentence Structure / Cấu Trúc Câu")
        style_guide.append(f"- **Average sentence length:** {avg_sentence_length:.1f} words")
        style_guide.append(f"- **Style:** {sentence_style}")
        style_guide.append("")
        
        if common_starters:
            style_guide.append("## Common Sentence Starters / Các Cách Mở Đầu Câu Thường Dùng")
            for starter, count in common_starters:
                if count > 1:
                    style_guide.append(f"- \"{starter}...\" (used {count}x)")
            style_guide.append("")
        
        style_guide.append("## Sample Sentences to Emulate / Các Câu Mẫu Để Tham Khảo")
        for i, sent in enumerate(sample_sentences[:5], 1):
            truncated = sent[:150] + "..." if len(sent) > 150 else sent
            style_guide.append(f"{i}. \"{truncated}\"")
        style_guide.append("")
        
        style_guide.append("---")
        style_guide.append("> **INSTRUCTION / HƯỚNG DẪN:**")
        style_guide.append("> - Write in the SAME formality level as this document")
        style_guide.append("> - Use similar sentence length and structure")
        style_guide.append("> - Reference the sample sentences for vocabulary and phrasing")
        style_guide.append("> - Maintain consistent tone throughout your output")
        
        return "\n".join(style_guide)

    def get_tools(self) -> List[FunctionTool]:
        return [
            FunctionTool(self.analyze_document_structure),
            FunctionTool(self.analyze_writing_style)
        ]
