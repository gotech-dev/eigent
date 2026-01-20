import os
import re
from typing import List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.utils.toolkit.abstract_toolkit import AbstractToolkit
from app.utils.listen.toolkit_listen import listen_toolkit
from app.service.task import Agents
from camel.toolkits.function_tool import FunctionTool

class DocxToolkit(AbstractToolkit):
    agent_name: str = Agents.document_agent

    def __init__(self, api_task_id: str, working_directory: str | None = None):
        self.api_task_id = api_task_id
        if working_directory is None:
            from app.component.environment import env
            working_directory = env("file_save_path", os.path.expanduser("~/Downloads"))
        self.working_directory = working_directory

    @listen_toolkit(
        None,
        lambda self, content, filename: f"Creating high-quality Word document: {filename}"
    )
    def create_professional_docx(self, content: str, filename: str) -> str:
        """
        Creates a professionally formatted Word document (.docx) from Markdown content.
        This tool ensures Times New Roman font, size 13, proper heading levels,
        bold text handling, and justified alignment with 1.3 line spacing.
        It also removes redundant Markdown characters like ** or ####.

        Args:
            content (str): The Markdown content to be converted into the Word document.
            filename (str): The name of the file to save (e.g., 'proposal.docx').

        Returns:
            str: A confirmation message with the absolute path of the created file.
        """
        if not filename.lower().endswith('.docx'):
            filename += '.docx'

        file_path = os.path.abspath(os.path.join(self.working_directory, filename))

        try:
            doc = Document()

            # Default style setup
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(13)

            # Administrative Margins (Standard VN)
            for section in doc.sections:
                section.top_margin = Inches(0.79)    # 2cm
                section.bottom_margin = Inches(0.79) # 2cm
                section.left_margin = Inches(1.18)   # 3cm
                section.right_margin = Inches(0.59)  # 1.5cm

            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # Header 1 (Title)
                if line.startswith('# '):
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run(line[2:].upper())
                    run.bold = True
                    run.font.size = Pt(16)

                # Header 2 (Chapter)
                elif line.startswith('## '):
                    p = doc.add_paragraph()
                    run = p.add_run(line[3:])
                    run.bold = True
                    run.font.size = Pt(14)

                # Header 3 (Section)
                elif line.startswith('### '):
                    text = line[4:].replace('**', '')
                    p = doc.add_paragraph()
                    run = p.add_run(text)
                    run.bold = True
                    run.font.size = Pt(13)

                # List items
                elif line.startswith('* ') or line.startswith('- '):
                    p = doc.add_paragraph(style='List Bullet')
                    text = line[2:]
                    self._add_formatted_text(p, text)

                # Regular text
                else:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.line_spacing = 1.3
                    self._add_formatted_text(p, line)

            doc.save(file_path)
            return f"Professional DOCX successfully created and saved at: {file_path}"
        except Exception as e:
            return f"Failed to create DOCX: {str(e)}"

    def _add_formatted_text(self, paragraph, text):
        """Helper to handle bold text within a paragraph"""
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def get_tools(self) -> List[FunctionTool]:
        return [FunctionTool(self.create_professional_docx)]
