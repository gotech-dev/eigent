
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os

def get_or_create_style(doc, style_name, style_type):
    try:
        return doc.styles[style_name]
    except KeyError:
        return doc.styles.add_style(style_name, style_type)

def copy_para_props_to_style(para, style):
    if para.runs:
        run = para.runs[0]
        style.font.name = run.font.name
        if run.font.size: style.font.size = run.font.size
        style.font.bold = True
        if run.font.color and run.font.color.rgb:
            style.font.color.rgb = run.font.color.rgb
    
    style.paragraph_format.alignment = para.alignment
    # Capture spacing if available
    if para.paragraph_format.space_before:
        style.paragraph_format.space_before = para.paragraph_format.space_before
    if para.paragraph_format.space_after:
        style.paragraph_format.space_after = para.paragraph_format.space_after

def standardize_reference_styles(input_path, output_path):
    doc = docx.Document(input_path)
    print(f"Standardizing styles for: {input_path}")
    
    # Identify "Title" (usually first centered block)
    for para in doc.paragraphs:
        if para.alignment == WD_ALIGN_PARAGRAPH.CENTER and para.text.strip():
            style = get_or_create_style(doc, 'Title', 1)
            copy_para_props_to_style(para, style)
            # Ensure Title is centered
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            print(f"Standardized 'Title' style from: '{para.text[:30]}...'")
            break

    # Identify "Heading 1" (Uppercase, Centered, Bold)
    for para in doc.paragraphs:
        if para.text.isupper() and para.alignment == WD_ALIGN_PARAGRAPH.CENTER and para.text.strip():
            style = get_or_create_style(doc, 'Heading 1', 1)
            copy_para_props_to_style(para, style)
            print(f"Standardized 'Heading 1' style from: '{para.text[:30]}...'")
            break

    # Identify "Heading 2" (Just Bold, maybe larger)
    for para in doc.paragraphs:
        if not para.text.isupper() and para.runs and para.runs[0].bold and para.text.strip():
             # Avoid re-using Title or H1
             if para.alignment != WD_ALIGN_PARAGRAPH.CENTER:
                style = get_or_create_style(doc, 'Heading 2', 1)
                copy_para_props_to_style(para, style)
                print(f"Standardized 'Heading 2' style from: '{para.text[:30]}...'")
                break

    # Identify "Normal" style (body text - for font consistency)
    for para in doc.paragraphs:
        if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY and para.text.strip():
            style = get_or_create_style(doc, 'Normal', 1)
            if para.runs:
                run = para.runs[0]
                style.font.name = run.font.name
                if run.font.size: style.font.size = run.font.size
            print(f"Standardized 'Normal' style from: '{para.text[:30]}...'")
            break

    doc.save(output_path)
    print(f"Saved standardized template to: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python standardize_styles.py <input.docx> <output.docx>")
    else:
        standardize_reference_styles(sys.argv[1], sys.argv[2])
